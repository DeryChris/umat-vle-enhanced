"""
Word document export for quiz questions.
Generates professional academic examination documents (.docx).

Supports three export types:
  - question_paper: Questions only (no answers)
  - answer_key: Answers, marks, and explanations
  - examiner_copy: Questions and answers combined

Supports versioned assessment papers (A/B/C) with randomized
question order and MC option order.
"""

import io
import logging
import random
import copy
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


# ── Version seeds ──────────────────────────────────────────

VERSION_SEEDS = {"A": 1001, "B": 2002, "C": 3003}


# ── Normalization ──────────────────────────────────────────

def _normalize_doc_settings(settings: Any) -> Dict[str, Any]:
    """
    Normalize doc_settings to ensure all nested fields have correct types.

    PHP encodes empty JS objects {} as [] (JSON array). This function
    detects and corrects such type mismatches at the API boundary.

    Raises:
        TypeError: If a critical field has an unfixable type mismatch.
    """
    if not isinstance(settings, dict):
        logger.warning("[EXPORT] doc_settings is not a dict: type=%s, converting to dict", type(settings).__name__)
        if isinstance(settings, list):
            return {}
        return {}

    # Fix student_info_fields: must be dict, not list
    sif = settings.get("student_info_fields")
    if sif is not None and not isinstance(sif, dict):
        logger.warning(
            "[EXPORT] student_info_fields has wrong type: expected=dict, actual=%s. Converting to default.",
            type(sif).__name__
        )
        if isinstance(sif, list) and len(sif) > 0:
            # If it's a non-empty list, something is very wrong
            logger.error("[EXPORT] student_info_fields is a non-empty list: %s. Using default.", _safe_summary(sif))
        settings["student_info_fields"] = {"studentName": True, "studentId": True}

    # Fix versions: must be list
    versions = settings.get("versions")
    if versions is not None:
        if isinstance(versions, str):
            settings["versions"] = [versions]
        elif not isinstance(versions, list):
            logger.warning("[EXPORT] versions has wrong type: %s, converting to list", type(versions).__name__)
            settings["versions"] = ["A"]

    # Fix boolean fields that might be wrong type
    bool_fields = ["show_page_numbers", "show_marks", "show_student_fields"]
    for field in bool_fields:
        val = settings.get(field)
        if val is not None and not isinstance(val, bool):
            if isinstance(val, str):
                settings[field] = val.lower() in ("true", "1", "yes")
            elif isinstance(val, (int, float)):
                settings[field] = bool(val)
            else:
                logger.warning("[EXPORT] %s has wrong type: %s, using True default", field, type(val).__name__)
                settings[field] = True

    # Fix numeric fields
    numeric_fields = ["duration", "total_marks", "marks_per_question", "answer_spaces"]
    for field in numeric_fields:
        val = settings.get(field)
        if val is not None and not isinstance(val, (int, float)):
            try:
                settings[field] = float(val) if "." in str(val) else int(val)
            except (ValueError, TypeError):
                logger.warning("[EXPORT] %s has non-numeric value: %s, using default", field, _safe_summary(val))
                defaults = {"duration": 120, "total_marks": 100, "marks_per_question": 1, "answer_spaces": 0}
                settings[field] = defaults.get(field, 0)

    return settings


# ── Logging helpers ────────────────────────────────────────

def _safe_summary(value: Any, max_len: int = 100) -> str:
    """Return a safe summary of a value for logging."""
    if value is None:
        return "None"
    t = type(value).__name__
    if isinstance(value, list):
        return f"list(len={len(value)})"
    if isinstance(value, dict):
        return f"dict(keys={list(value.keys())[:5]})"
    s = str(value)
    if len(s) > max_len:
        s = s[:max_len] + "..."
    return f"{t}({s})"


def _log_export_payload(
    questions: List[Dict],
    doc_settings: Dict,
    export_type: str,
    version: str,
) -> None:
    """Log structured export payload for debugging."""
    logger.info("[EXPORT] Starting document generation")
    logger.info("[EXPORT]   export_type=%s  version=%s", export_type, version)
    logger.info("[EXPORT]   questions: %s", _safe_summary(questions))
    logger.info("[EXPORT]   doc_settings keys: %s", list(doc_settings.keys()) if isinstance(doc_settings, dict) else _safe_summary(doc_settings))

    # Log critical nested fields
    critical_fields = [
        "student_info_fields", "versions", "answer_spaces",
        "marks_per_question", "show_marks", "show_page_numbers",
    ]
    for field in critical_fields:
        if isinstance(doc_settings, dict) and field in doc_settings:
            val = doc_settings[field]
            logger.info("[EXPORT]   doc_settings.%s: type=%s value=%s", field, type(val).__name__, _safe_summary(val))

    # Log each question's structure
    for i, q in enumerate(questions[:3]):  # First 3 only
        if isinstance(q, dict):
            logger.info("[EXPORT]   question[%d]: type=%s keys=%s", i, type(q).__name__, list(q.keys()))
            # Check options specifically
            opts = q.get("options")
            if opts is not None:
                logger.info("[EXPORT]   question[%d].options: type=%s len=%s", i, type(opts).__name__, len(opts) if hasattr(opts, '__len__') else "N/A")
        else:
            logger.warning("[EXPORT]   question[%d]: UNEXPECTED TYPE=%s value=%s", i, type(q).__name__, _safe_summary(q))


# ── Public API ─────────────────────────────────────────────

def generate_document(
    questions: List[Dict[str, Any]],
    export_type: str,
    doc_settings: Dict[str, Any],
    version: str = "A",
) -> bytes:
    """
    Generate a .docx document and return its bytes.

    Args:
        questions:     List of question dicts from AI generation.
        export_type:   "question_paper" | "answer_key" | "examiner_copy"
        doc_settings:  Document metadata and layout settings.
        version:       "A", "B", or "C".

    Returns:
        Raw bytes of the .docx file.
    """
    # Structured logging before generation
    _log_export_payload(questions, doc_settings, export_type, version)

    # Normalize doc_settings at boundary
    doc_settings = _normalize_doc_settings(doc_settings)

    rng = random.Random(VERSION_SEEDS.get(version, 1001))

    ordered = _apply_version_order(questions, version, rng)
    sections = _group_by_type(ordered)

    doc = Document()
    _setup_page(doc, doc_settings)
    _add_header(doc, doc_settings, version, export_type)
    _add_student_info(doc, doc_settings)
    _add_instructions(doc, doc_settings, export_type)
    _add_questions(doc, sections, export_type, doc_settings, rng)
    _add_footer(doc, doc_settings)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Version ordering ───────────────────────────────────────

def _apply_version_order(
    questions: List[Dict], version: str, rng: random.Random
) -> List[Dict]:
    qs = copy.deepcopy(questions)
    rng.shuffle(qs)

    for q in qs:
        if q.get("type") == "multichoice" and q.get("options"):
            correct_text = None
            idx = q.get("correct_answer_index", 0)
            if 0 <= idx < len(q["options"]):
                correct_text = q["options"][idx]

            pairs = list(enumerate(q["options"]))
            rng.shuffle(pairs)
            new_options = [p[1] for p in pairs]
            q["options"] = new_options

            if correct_text is not None:
                for i, opt in enumerate(new_options):
                    if opt == correct_text:
                        q["correct_answer_index"] = i
                        break

    return qs


# ── Grouping ───────────────────────────────────────────────

def _group_by_type(questions: List[Dict]) -> Dict[str, List[Dict]]:
    groups: Dict[str, List[Dict]] = {}
    for q in questions:
        t = q.get("type", "multichoice")
        groups.setdefault(t, []).append(q)
    return groups


# ── Page setup ─────────────────────────────────────────────

def _setup_page(doc: Document, settings: Dict):
    section = doc.sections[0]

    orientation = settings.get("orientation", "portrait")
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        new_w, new_h = section.page_height, section.page_width
        section.page_width = new_w
        section.page_height = new_h
    else:
        section.orientation = WD_ORIENT.PORTRAIT

    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15


# ── Header ─────────────────────────────────────────────────

def _add_header(doc: Document, settings: Dict, version: str, export_type: str):
    institution = settings.get("institution_name", "")
    if institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(institution)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

    title = settings.get("assessment_title", "Assessment")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    version_label = f" — Version {version}" if version and version != "A" else ""
    export_label = {
        "question_paper": "Question Paper",
        "answer_key": "Answer Key / Marking Scheme",
        "examiner_copy": "Examiner's Copy",
    }.get(export_type, "")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(export_label + version_label)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    course_title = settings.get("course_title", "")
    course_code = settings.get("course_code", "")
    dept = settings.get("department", "")
    lecturer = settings.get("lecturer_name", "")

    meta_parts = []
    if course_code or course_title:
        label = f"{course_code} — {course_title}" if course_code and course_title else (course_code or course_title)
        meta_parts.append(label)
    if dept:
        meta_parts.append(dept)
    if lecturer:
        meta_parts.append(f"Lecturer: {lecturer}")

    if meta_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(meta_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)

    date_str = settings.get("examination_date_display", "") or settings.get("examination_date", "")
    duration = settings.get("duration", 0)
    detail_parts = []
    if date_str:
        detail_parts.append(date_str)
    if duration:
        detail_parts.append(f"Duration: {duration} Minutes")
    total_marks = settings.get("total_marks", 0)
    if total_marks:
        detail_parts.append(f"Total Marks: {total_marks}")

    if detail_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(detail_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)

    _add_thin_line(doc)


def _add_thin_line(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)


# ── Student info block ─────────────────────────────────────

def _add_student_info(doc: Document, settings: Dict):
    fields_cfg = settings.get("student_info_fields", None)

    # Guard: PHP encodes empty JS objects as [] (list), not {} (dict)
    # Also handle any non-dict type that might slip through
    if not isinstance(fields_cfg, dict):
        if not settings.get("show_student_fields", True):
            return
        fields_cfg = {"studentName": True, "studentId": True}

    info_fields = []
    if fields_cfg.get("studentName"):
        info_fields.append(("Name:", "_" * 40))
    if fields_cfg.get("studentId"):
        info_fields.append(("Index Number:", "_" * 40))
    if fields_cfg.get("class"):
        info_fields.append(("Class:", "_" * 40))
    if fields_cfg.get("programme"):
        info_fields.append(("Programme:", "_" * 40))
    if fields_cfg.get("level"):
        info_fields.append(("Level:", "_" * 40))
    if fields_cfg.get("signature"):
        info_fields.append(("Signature:", "_" * 40))

    if not info_fields:
        return

    table = doc.add_table(rows=len(info_fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, line) in enumerate(info_fields):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Cm(3.5)
        c1.width = Cm(12)

        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(11)
        r1 = c1.paragraphs[0].add_run(line)
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(150, 150, 150)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    doc.add_paragraph("")


# ── Instructions ───────────────────────────────────────────

def _add_instructions(doc: Document, settings: Dict, export_type: str):
    if export_type != "question_paper":
        return

    custom = settings.get("candidate_instructions", "").strip()
    if custom:
        p = doc.add_paragraph()
        run = p.add_run("Instructions:")
        run.bold = True
        run.font.size = Pt(11)

        for line in custom.split("\n"):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line.strip())
            run.font.size = Pt(10)

        doc.add_paragraph("")


# ── Questions ──────────────────────────────────────────────

SECTION_LABELS = {
    "multichoice": "Multiple Choice Questions",
    "truefalse": "True / False Questions",
    "shortanswer": "Short Answer Questions",
}


def _add_questions(
    doc: Document,
    sections: Dict[str, List[Dict]],
    export_type: str,
    settings: Dict,
    rng: random.Random,
):
    global_num = 0
    marks_per_q = settings.get("marks_per_question", 1)
    show_marks = settings.get("show_marks", True)
    answer_spaces = settings.get("answer_spaces", 0)

    for qtype, qs in sections.items():
        sec_title = SECTION_LABELS.get(qtype, qtype.title())
        p = doc.add_paragraph()
        run = p.add_run(sec_title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)
        _add_thin_line(doc)

        for q in qs:
            global_num += 1
            marks = q.get("marks", marks_per_q)

            _add_question_body(doc, global_num, q, marks, show_marks, export_type, answer_spaces)

            if export_type in ("answer_key", "examiner_copy"):
                _add_answer_block(doc, q, marks)

            doc.add_paragraph("")


def _add_question_body(
    doc: Document,
    num: int,
    q: Dict,
    marks: float,
    show_marks: bool,
    export_type: str,
    answer_spaces: int,
):
    p = doc.add_paragraph()
    qtype = q.get("type", "")

    marks_str = f"  [{marks} mark{'s' if marks != 1 else ''}]" if show_marks else ""

    run = p.add_run(f"{num}. ")
    run.bold = True
    run.font.size = Pt(12)

    text = q.get("question_text", "")
    run = p.add_run(text)
    run.font.size = Pt(12)

    if marks_str:
        run = p.add_run(marks_str)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

    options = q.get("options") or []
    if options:
        for i, opt in enumerate(options):
            letter = chr(65 + i)
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Cm(1.0)
            run = op.add_run(f"{letter}.  {opt}")
            run.font.size = Pt(11)

    if qtype == "shortanswer" and export_type == "question_paper" and answer_spaces > 0:
        spaces = max(answer_spaces, 1)
        for _ in range(spaces):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Cm(0.5)
            run = sp.add_run("_" * 80)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 180, 180)


def _add_answer_block(doc: Document, q: Dict, marks: float):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)

    if q.get("type") == "shortanswer":
        answer = q.get("correct_text", "N/A")
    else:
        idx = q.get("correct_answer_index", 0)
        opts = q.get("options", [])
        if 0 <= idx < len(opts):
            answer = f"{chr(65 + idx)}. {opts[idx]}"
        else:
            answer = "N/A"

    run = p.add_run("Answer: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 107, 47)
    run = p.add_run(answer)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 107, 47)

    fb_correct = q.get("feedback_correct", "").strip()
    fb_incorrect = q.get("feedback_incorrect", "").strip()
    explanation = fb_correct or fb_incorrect
    if explanation:
        ep = doc.add_paragraph()
        ep.paragraph_format.left_indent = Cm(1.0)
        run = ep.add_run("Explanation: ")
        run.bold = True
        run.font.size = Pt(10)
        run = ep.add_run(explanation)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)


# ── Total marks ────────────────────────────────────────────

def _add_total_marks(doc: Document, questions: List[Dict], marks_per_q: float):
    total = sum(q.get("marks", marks_per_q) for q in questions)
    p = doc.add_paragraph()
    _add_thin_line(doc)
    run = p.add_run(f"TOTAL MARKS: {total}")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ── Footer ─────────────────────────────────────────────────

def _add_footer(doc: Document, settings: Dict):
    show_pn = settings.get("show_page_numbers", True)
    if not show_pn:
        return

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar2)

    run2 = p.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2._r.append(fldChar3)

    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run2._r.append(instrText2)

    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2._r.append(fldChar4)
